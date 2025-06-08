from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import re

def printList(keys, values, save_path="C:\\Users\\HP\\Desktop\\Note\\"):
    doc = Document()
    doc.add_heading('Resume', 0)
    section = doc.sections[-1]
    section.page_width = Pt(595.3)  
    section.page_height = Pt(841.9)  


    section._sectPr.xpath('./w:cols')[0].set('num', '2')
   
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    def add_section_heading(title):
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(14)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        doc.add_paragraph("-" *40)

    def add_field(label, value):
        if value and value.strip().lower() not in ['none', 'null', '']:
            para = doc.add_paragraph()
            para.add_run(f"{label}: ").bold = True
            para.add_run(value.strip())

    def add_multi_paragraph_field(label, text):
        if text and text.strip().lower() not in ['none', 'null', '']:
            doc.add_paragraph(f"{label}:", style='Normal').runs[0].bold = True
            project_paragraphs = re.split(r'\.\s*(?=[A-Z])', text.strip())
            for para in project_paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip() + '.', style='Normal')

    data = dict(zip(keys, values))

    add_section_heading("1. Personal Information")
    for label in ["Full Name", "Email", "Contact", "City", "Occupation"]:
        add_field(label, data.get(label))

    add_section_heading("2. College and Branch Info")
    for label in ["Degree", "College", "CGPA", "Current Year", "Branch"]:
        add_field(label, data.get(label))

    add_section_heading("3. Technical Skills")
    for label in ["Language (comma)", "Tools [comma separate]", "Frameworks"]:
        add_field(label, data.get(label))

    add_section_heading("4. Projects")
    add_field("Tools/Tech Used", data.get("Tools/Tech Used"))
    add_multi_paragraph_field("Short description (Project)", data.get("Short description (Project)"))
    add_field("GitHub link (if any)", data.get("GitHub link (if any)"))

    add_section_heading("5. Certifications and Awards")
    add_field("Certification", data.get("Certification"))
    add_field("Awards (If Any)", data.get("Awards (If Any)"))

    add_section_heading("6. Experience and Hobbies")
    add_field("Experience", data.get("Experience"))
    add_field("Interests", data.get("Interests"))
    add_field("Intern (If Any)", data.get("Intern (If Any)"))

   
    filename = data.get("Full Name", "ATS_Resume").strip().replace(" ", "_")
    if not filename.endswith(".docx"):
        filename += ".docx"

    os.makedirs(save_path, exist_ok=True)
    full_path = os.path.join(save_path, filename)
    doc.save(full_path)
    print(f"✅ Resume saved at: {full_path}")
